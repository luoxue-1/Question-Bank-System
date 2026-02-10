from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER

class Exporter:
    """导出功能类，负责将试卷和答案导出为Word和PDF格式"""
    
    def export_paper_to_word(self, paper_questions, filename='试卷.docx'):
        """
        导出试卷为Word文件
        
        Args:
            paper_questions: 试卷题目列表
            filename: 导出文件名
        """
        doc = Document()
        
        # 添加标题
        doc.add_heading('试卷', 0)
        
        # 添加题目
        for q in paper_questions:
            doc.add_heading(f"{q['id']}. {q['topic']}", level=2)
            doc.add_paragraph(q['content'])
            doc.add_paragraph('')  # 空行
        
        # 保存文件
        doc.save(filename)
    
    def export_answers_to_word(self, answers, filename='答案.docx'):
        """
        导出答案为Word文件
        
        Args:
            answers: 答案列表
            filename: 导出文件名
        """
        doc = Document()
        
        # 添加标题
        doc.add_heading('答案', 0)
        
        # 添加答案
        for ans in answers:
            doc.add_heading(f"{ans['id']}. {ans['topic']}", level=2)
            doc.add_paragraph(f"题目: {ans['content']}")
            doc.add_paragraph(f"答案: {ans['answer']}")
            doc.add_paragraph('')  # 空行
        
        # 保存文件
        doc.save(filename)
    
    def export_paper_to_pdf(self, paper_questions, filename='试卷.pdf'):
        """
        导出试卷为PDF文件
        
        Args:
            paper_questions: 试卷题目列表
            filename: 导出文件名
        """
        doc = SimpleDocTemplate(filename, pagesize=A4)
        styles = getSampleStyleSheet()
        
        # 创建自定义样式
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Title'],
            alignment=TA_CENTER,
            spaceAfter=30
        )
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading2'],
            spaceAfter=12
        )
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            alignment=TA_LEFT,
            spaceAfter=18
        )
        
        # 构建内容
        content = []
        
        # 添加标题
        content.append(Paragraph('试卷', title_style))
        
        # 添加题目
        for q in paper_questions:
            content.append(Paragraph(f"{q['id']}. {q['topic']}", heading_style))
            content.append(Paragraph(q['content'], normal_style))
            content.append(Spacer(1, 12))
        
        # 生成PDF
        doc.build(content)
    
    def export_answers_to_pdf(self, answers, filename='答案.pdf'):
        """
        导出答案为PDF文件
        
        Args:
            answers: 答案列表
            filename: 导出文件名
        """
        doc = SimpleDocTemplate(filename, pagesize=A4)
        styles = getSampleStyleSheet()
        
        # 创建自定义样式
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Title'],
            alignment=TA_CENTER,
            spaceAfter=30
        )
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading2'],
            spaceAfter=12
        )
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            alignment=TA_LEFT,
            spaceAfter=12
        )
        
        # 构建内容
        content = []
        
        # 添加标题
        content.append(Paragraph('答案', title_style))
        
        # 添加答案
        for ans in answers:
            content.append(Paragraph(f"{ans['id']}. {ans['topic']}", heading_style))
            content.append(Paragraph(f"题目: {ans['content']}", normal_style))
            content.append(Paragraph(f"答案: {ans['answer']}", normal_style))
            content.append(Spacer(1, 18))
        
        # 生成PDF
        doc.build(content)
